from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUTPUT_DIR / "poyasnitelnaya_zapiska.docx"
MD_PATH = OUTPUT_DIR / "poyasnitelnaya_zapiska.md"
README_PATH = OUTPUT_DIR / "README.md"


PROJECT_TITLE = "Визуализатор алгоритма Дейкстры"
COURSEWORK_TITLE = (
    "Программная реализация алгоритма нахождения кратчайших путей "
    "в графах методом Дейкстры"
)


# ---------------------------------------------------------------------------
# Низкоуровневые помощники оформления
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_paragraph(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    return p


def add_heading(doc, title, level=1):
    p = doc.add_paragraph()
    # Требование: перед заголовком интервал примерно в 2 раза,
    # после — в 1,5 раза больше расстояния между строками текста.
    p.paragraph_format.space_before = Pt(24 if level == 1 else 18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.keep_with_next = True
    # Требование: заголовки прописными буквами, жирным, по центру, без переносов.
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


_TABLE_COUNTER = {"n": 0}


def add_table(doc, caption, headers, rows, widths=None):
    # Требование: таблицы нумеруются и подписываются, в тексте на них есть ссылки.
    _TABLE_COUNTER["n"] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(4)
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_run = cap.add_run(f"Таблица {_TABLE_COUNTER['n']} — {caption}")
    cap_run.font.name = "Times New Roman"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cap_run.font.size = Pt(14)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr_cells[idx], header, bold=True, size=11)
        set_cell_shading(hdr_cells[idx], "E8EEF5")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=11)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)


def configure_document(doc):
    section = doc.sections[0]
    # Поля по требованию КР: верхнее 1,5 / нижнее 1,5 / левое 2,5 / правое 1,0 см. A4.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    # Требование: нумерация листов — в правом верхнем углу (в верхнем колонтитуле).
    header = section.header.paragraphs[0]
    add_page_number(header, alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_title_page(doc):
    for text, size, bold in [
        ("[Название образовательной организации]", 12, False),
        ("[Факультет / кафедра]", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА\nк курсовой работе")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'по теме: «{COURSEWORK_TITLE}»')
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)

    for _ in range(5):
        doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    rows = [
        ("Выполнил:", "[ФИО студента]"),
        ("Группа:", "[номер группы]"),
        ("Проверил:", "[ФИО руководителя]"),
        ("Оценка:", "____________"),
    ]
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            set_cell_text(cell, value, bold=(col_idx == 0), size=12)

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Город] 2026")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    doc.add_page_break()


def add_contents_page(doc):
    add_heading(doc, "Содержание", 1)
    add_numbered(doc, [
        "Титульный лист.",
        "Введение.",
        "Краткое описание используемых методов и алгоритмов.",
        "Особенности программной реализации алгоритмов.",
        "Сравнительный анализ используемых алгоритмов и методов.",
        "Листинги программ с обязательными исчерпывающими комментариями.",
        "Распечатки содержимого входного и выходного файлов программы.",
        "Список использованных литературных и информационных источников.",
    ])
    add_paragraph(
        doc,
        "Структура пояснительной записки соответствует требованиям к курсовой работе: "
        "сначала формулируется задача, затем описываются методы и алгоритмы, особенности "
        "программной реализации, проводится сравнительный анализ, приводятся листинги "
        "с комментариями и распечатки входного и выходного файлов программы."
    )
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Основная сборка DOCX
# ---------------------------------------------------------------------------
def build_docx():
    doc = Document()
    _TABLE_COUNTER["n"] = 0
    configure_document(doc)
    add_title_page(doc)
    add_contents_page(doc)

    # ===================================================================
    # 2. ВВЕДЕНИЕ
    # ===================================================================
    add_heading(doc, "2. Введение", 1)
    add_paragraph(
        doc,
        "В курсовой работе рассматривается задача поиска кратчайшего пути во взвешенном "
        "графе — нахождение маршрута наименьшей суммарной стоимости между двумя вершинами "
        "при неотрицательных весах рёбер. Эта задача широко встречается в маршрутизации "
        "компьютерных сетей, навигации, логистике и планировании."
    )
    add_paragraph(
        doc,
        "Известными методами её решения являются поиск в ширину (для невзвешенных графов), "
        "алгоритм Дейкстры, алгоритм Беллмана-Форда, алгоритм Флойда-Уоршелла и алгоритм A*. "
        "По заданию основным выбран алгоритм Дейкстры, так как он гарантирует нахождение "
        "оптимального пути в графах с неотрицательными весами и обеспечивает хорошую "
        "временную сложность при использовании очереди с приоритетами."
    )
    add_paragraph(
        doc,
        "Целью работы является самостоятельная программная реализация алгоритма Дейкстры "
        "и сравнительный анализ его с другими методами поиска путей. Алгоритм и очередь "
        "с приоритетами реализованы вручную, без использования готовых библиотечных "
        "структур. Для наглядной демонстрации работы создано интерактивное приложение на "
        "HTML, CSS и JavaScript; библиотеки Cytoscape.js, Chart.js и jsPDF применяются "
        "только для отображения графа, построения графиков и формирования отчёта."
    )

    # ===================================================================
    # 3. КРАТКОЕ ОПИСАНИЕ МЕТОДОВ И АЛГОРИТМОВ
    # ===================================================================
    add_heading(doc, "3. Краткое описание используемых методов и алгоритмов", 1)
    add_heading(doc, "3.1. Графовая модель", 2)
    add_paragraph(
        doc,
        "Граф задаётся множеством вершин V и множеством рёбер E. В проекте поддерживаются "
        "как неориентированные, так и ориентированные графы. Каждое ребро имеет вес — "
        "неотрицательную числовую стоимость перехода между вершинами. Основные понятия "
        "графовой модели и их назначение в программе приведены в таблице 1."
    )
    add_table(
        doc,
        "Основные понятия графовой модели",
        ["Понятие", "Назначение в программе"],
        [
            ["Вершина", "Объект с идентификатором, подписью и координатами на холсте."],
            ["Ребро", "Связь между двумя вершинами с неотрицательным весом."],
            ["Список смежности", "Способ быстрого получения соседей выбранной вершины."],
            ["Очередь с приоритетами", "Структура для выбора вершины с минимальным расстоянием."],
        ],
        [3.5, 11.0],
    )
    add_heading(doc, "3.2. Алгоритм Дейкстры", 2)
    add_paragraph(
        doc,
        "Алгоритм Дейкстры находит кратчайшие пути от стартовой вершины до остальных "
        "вершин графа. На каждом шаге выбирается ещё не обработанная вершина с минимальным "
        "известным расстоянием, после чего выполняется релаксация её исходящих рёбер. "
        "Основные шаги алгоритма:"
    )
    add_numbered(doc, [
        "Для стартовой вершины устанавливается расстояние 0, для остальных — бесконечность.",
        "Стартовая вершина помещается в очередь с приоритетом 0.",
        "Из очереди извлекается вершина с минимальным расстоянием.",
        "Для каждого соседа проверяется, можно ли улучшить известное расстояние (релаксация).",
        "Если путь улучшен, обновляются расстояние и предшественник вершины.",
        "После обработки всех достижимых вершин путь восстанавливается по предшественникам.",
    ])
    add_paragraph(
        doc,
        "Очередь с приоритетами реализована на основе бинарной кучи (min-heap), что "
        "позволяет выполнять вставку и извлечение минимального элемента за время O(log n). "
        "Общая временная сложность алгоритма составляет O((V + E) log V), а затраты памяти — "
        "O(V + E) на хранение графа, расстояний, предшественников и очереди."
    )

    # ===================================================================
    # 4. ОСОБЕННОСТИ ПРОГРАММНОЙ РЕАЛИЗАЦИИ
    # ===================================================================
    add_heading(doc, "4. Особенности программной реализации алгоритмов", 1)
    add_heading(doc, "4.1. Общая структура проекта", 2)
    add_paragraph(
        doc,
        "Программа разделена на независимые модули, каждый из которых отвечает за свою "
        "часть функциональности. Такое разделение упрощает отладку и повторное "
        "использование кода. Назначение основных файлов проекта приведено в таблице 2."
    )
    add_table(
        doc,
        "Назначение основных файлов проекта",
        ["Файл", "Назначение"],
        [
            ["index.html", "Разметка интерфейса и подключение локальных библиотек."],
            ["css/style.css", "Темы, адаптивность, стили графа и элементов интерфейса."],
            ["js/graph.js", "Структура графа, генерация, импорт, экспорт и валидация данных."],
            ["js/dijkstra.js", "Алгоритм Дейкстры и очередь с приоритетами (ядро проекта)."],
            ["js/visualizer.js", "Пошаговая визуализация состояния алгоритма."],
            ["js/ui.js", "Главный контроллер интерфейса и обработчики событий."],
            ["js/storage.js", "LocalStorage, импорт/экспорт, формирование PNG и PDF."],
            ["tests/unit-tests.js", "Автоматические тесты ядра проекта."],
        ],
        [4.2, 10.3],
    )
    add_paragraph(
        doc,
        "Общая логика программной реализации алгоритма Дейкстры строго соответствует "
        "классической схеме (см. блок-схему алгоритма). На схеме отражены основные этапы: "
        "инициализация структур данных, цикл извлечения вершины с минимальным приоритетом "
        "из очереди, релаксация смежных рёбер и восстановление итогового кратчайшего пути."
    )
    add_heading(doc, "4.2. Используемые типы данных", 2)
    add_paragraph(
        doc,
        "Все действия над данными выполняются над несколькими основными типами. Граф "
        "хранится в виде хеш-таблиц вершин и рёбер, что обеспечивает доступ к элементу за "
        "время O(1). Расстояния и предшественники хранятся в ассоциативных массивах Map, "
        "а множество обработанных вершин — в Set. Перечень основных типов данных и их "
        "назначение приведены в таблице 3."
    )
    add_table(
        doc,
        "Основные типы данных программной реализации",
        ["Тип данных", "Назначение"],
        [
            ["GraphManager", "Хранит вершины и рёбра, строит список смежности, проверяет данные."],
            ["DijkstraAlgorithm", "Реализует алгоритм; хранит distances, predecessors, visited."],
            ["PriorityQueue", "Бинарная мин-куча для выбора вершины с минимальным расстоянием."],
            ["Map<string, number>", "Таблица кратчайших расстояний distances."],
            ["Map<string, string>", "Массив предшественников predecessors для восстановления пути."],
            ["Set<string>", "Множество уже обработанных (посещённых) вершин visited."],
        ],
        [4.8, 9.7],
    )
    add_heading(doc, "4.3. Входные и выходные данные", 2)
    add_paragraph(
        doc,
        "По требованию к курсовой работе ввод-вывод данных организован посредством файлов. "
        "Входные данные представлены JSON-файлом графа: он содержит тип графа (поле "
        "directed), массив вершин nodes и массив рёбер edges. Для каждой вершины задаются "
        "идентификатор, подпись и координаты, для каждого ребра — источник, цель и вес. "
        "Загрузка выполняется через стандартный File API (чтение файла объектом FileReader)."
    )
    add_paragraph(
        doc,
        "Выходные данные сохраняются в файлы: граф экспортируется в JSON, результат работы "
        "алгоритма — в текстовый (TXT) и JSON-файлы, изображение графа — в PNG, а итоговый "
        "отчёт — в PDF. Результат включает найденный путь, его стоимость, таблицу "
        "расстояний, массив предшественников, время выполнения и число операций. Структура "
        "входного и выходного файлов подробно показана в разделе 7."
    )
    add_heading(doc, "4.4. Проверка корректности данных", 2)
    add_paragraph(
        doc,
        "Перед загрузкой граф проходит строгую проверку в методе "
        "GraphManager.validateGraphData: контролируются наличие массивов nodes и edges, "
        "уникальность идентификаторов вершин, существование вершин, на которые ссылаются "
        "рёбра, конечность и неотрицательность весов, отсутствие петель и дублирующихся "
        "рёбер. При нарушении формируется понятное сообщение об ошибке, а некорректные "
        "данные не загружаются."
    )
    add_heading(doc, "4.5. Набор тестов для проверки работоспособности", 2)
    add_paragraph(
        doc,
        "Для проверки работоспособности и отладки разработан набор автоматических тестов "
        "(файл tests/unit-tests.js), запускаемых командой npm test. Тесты проверяют:"
    )
    add_bullets(doc, [
        "корректность кратчайшего пути в неориентированном графе;",
        "обработку случая, когда путь между вершинами отсутствует;",
        "ошибки при несуществующей стартовой или конечной вершине;",
        "отказ алгоритма Дейкстры при наличии отрицательных весов;",
        "работу быстрого режима без накопления шагов визуализации;",
        "строгую проверку JSON: дубли идентификаторов, ссылки на несуществующие вершины, "
        "некорректные веса;",
        "корректность алгоритмов Беллмана-Форда и Флойда-Уоршелла, включая обнаружение "
        "отрицательного цикла.",
    ])
    add_heading(doc, "4.6. Ограничения производительности", 2)
    add_paragraph(
        doc,
        "Подробная пошаговая визуализация хранит снимок состояния алгоритма на каждом "
        "шаге, поэтому для больших графов она требует много памяти. По этой причине для "
        "графов более 100 вершин интерфейс рекомендует режим быстрого расчёта, а для "
        "графов более 200 вершин подробная визуализация отключается автоматически. "
        "Бенчмарки выполняют алгоритм без накопления шагов, что позволяет тестировать "
        "графы в 500–1000 вершин без переполнения памяти."
    )

    # ===================================================================
    # 5. СРАВНИТЕЛЬНЫЙ АНАЛИЗ
    # ===================================================================
    add_heading(doc, "5. Сравнительный анализ используемых алгоритмов и методов", 1)
    add_paragraph(
        doc,
        "Сравнительные оценки рассмотренных алгоритмов поиска путей по времени выполнения "
        "и объёму используемой памяти приведены в таблице 4. Здесь V — число вершин, "
        "E — число рёбер графа."
    )
    add_table(
        doc,
        "Сравнение алгоритмов поиска кратчайших путей",
        ["Метод", "Ограничения", "Время", "Память"],
        [
            ["Поиск в ширину (BFS)", "Только невзвешенные графы", "O(V + E)", "O(V)"],
            ["Дейкстра (бинарная куча)", "Нет отрицательных весов", "O((V + E) log V)", "O(V + E)"],
            ["Беллман-Форд", "Допускает отрицательные веса", "O(V·E)", "O(V)"],
            ["Флойд-Уоршелл", "Все пары вершин", "O(V³)", "O(V²)"],
            ["A*", "Требует эвристику", "Зависит от эвристики", "O(V)"],
        ],
        [4.0, 4.0, 3.6, 2.9],
    )
    add_paragraph(
        doc,
        "Поиск в ширину работает быстрее всех, но применим лишь к невзвешенным графам. "
        "Беллман-Форд универсальнее Дейкстры (допускает отрицательные веса), но медленнее "
        "из-за оценки O(V·E). Флойд-Уоршелл удобен, когда нужны пути между всеми парами "
        "вершин, однако требует O(V²) памяти и O(V³) времени. Алгоритм A* эффективен при "
        "наличии хорошей эвристики, например на картах."
    )
    add_paragraph(
        doc,
        "Для данной курсовой работы выбран алгоритм Дейкстры, так как задача ограничена "
        "неотрицательными весами, а его сложность O((V + E) log V) при умеренных затратах "
        "памяти O(V + E) делает его оптимальным для поиска путей от одной стартовой вершины "
        "и наглядной пошаговой демонстрации."
    )

    # ===================================================================
    # 6. ЛИСТИНГИ
    # ===================================================================
    add_heading(doc, "6. Листинги программ с комментариями", 1)
    add_paragraph(
        doc,
        "Ниже приведены ключевые фрагменты исходного кода с комментариями. Полные тексты "
        "модулей находятся в папке js проекта."
    )
    add_heading(doc, "6.1. Очередь с приоритетами (бинарная куча)", 2)
    add_code_block(doc, """
class PriorityQueue {
    constructor() {
        // Массив-хранилище бинарной мин-кучи; элемент: { element, priority }
        this.heap = [];
    }

    // Вставка элемента за O(log n): новый элемент «всплывает» вверх кучи
    enqueue(element, priority) {
        this.heap.push({ element, priority });
        this._bubbleUp(this.heap.length - 1);
    }

    // Извлечение минимума (корня кучи) за O(log n)
    dequeue() {
        if (this.isEmpty()) return null;
        const min = this.heap[0];
        const last = this.heap.pop();
        if (this.heap.length > 0) {
            this.heap[0] = last;   // последний элемент перемещаем в корень
            this._sinkDown(0);     // и «проталкиваем» его вниз на своё место
        }
        return min;
    }

    isEmpty() { return this.heap.length === 0; }

    // «Всплытие»: меняем элемент с родителем, пока его приоритет меньше
    _bubbleUp(i) {
        while (i > 0) {
            const parent = Math.floor((i - 1) / 2);
            if (this.heap[i].priority >= this.heap[parent].priority) break;
            [this.heap[i], this.heap[parent]] = [this.heap[parent], this.heap[i]];
            i = parent;
        }
    }
}
""")
    add_heading(doc, "6.2. Основной цикл алгоритма Дейкстры", 2)
    add_code_block(doc, """
run(startNodeId, endNodeId) {
    const startTime = performance.now();

    // Сброс состояния перед запуском
    this.distances = new Map();      // кратчайшие расстояния
    this.predecessors = new Map();   // предшественники для восстановления пути
    this.visited = new Set();        // обработанные вершины
    this.operationCount = 0;

    this._validateInput(startNodeId, endNodeId); // проверка вершин
    this._validateWeights();                      // запрет отрицательных весов

    const adjList = this.graph.getAdjacencyList();
    const allNodes = Array.from(this.graph.nodes.keys());

    // Шаг 1. Инициализация: старт = 0, остальные = бесконечность
    allNodes.forEach(id => {
        this.distances.set(id, id === startNodeId ? 0 : Infinity);
        this.predecessors.set(id, null);
    });

    // Шаг 2. Стартовая вершина — в очередь с приоритетом 0
    const pq = new PriorityQueue();
    pq.enqueue(startNodeId, 0);

    // Шаг 3. Основной цикл
    while (!pq.isEmpty()) {
        const { element: u } = pq.dequeue();   // вершина с мин. расстоянием
        if (this.visited.has(u)) continue;     // пропуск устаревших записей
        this.visited.add(u);
        if (endNodeId && u === endNodeId) break; // ранний выход при достижении цели

        // Релаксация исходящих рёбер вершины u
        for (const { nodeId: v, weight } of (adjList.get(u) || [])) {
            if (this.visited.has(v)) continue;
            const newDist = this.distances.get(u) + weight;
            if (newDist < this.distances.get(v)) { // найден более короткий путь
                this.distances.set(v, newDist);
                this.predecessors.set(v, u);
                pq.enqueue(v, newDist);
            }
        }
    }

    // Шаг 4. Восстановление пути по массиву предшественников
    const path = this._reconstructPath(startNodeId, endNodeId);
    const executionTime = performance.now() - startTime;
    return { path, distances: this.distances, predecessors: this.predecessors,
             operationCount: this.operationCount, executionTime };
}
""")
    add_heading(doc, "6.3. Проверка импортируемого графа", 2)
    add_code_block(doc, """
static validateGraphData(data) {
    // 1) Корневая структура должна быть объектом с массивами nodes и edges
    if (!data || typeof data !== 'object')
        throw new Error('Данные графа должны быть объектом');
    if (!Array.isArray(data.nodes))
        throw new Error('Отсутствует массив nodes');
    if (!Array.isArray(data.edges))
        throw new Error('Отсутствует массив edges');

    // 2) Вершины: строковый уникальный id
    const nodeIds = new Set();
    data.nodes.forEach((node, i) => {
        if (typeof node.id !== 'string' || node.id.trim() === '')
            throw new Error(`Вершина #${i + 1}: отсутствует строковый id`);
        if (nodeIds.has(node.id))
            throw new Error(`Дублирующийся id вершины ${node.id}`);
        nodeIds.add(node.id);
    });

    // 3) Рёбра: ссылки на существующие вершины и конечный вес
    data.edges.forEach((edge) => {
        if (!nodeIds.has(edge.source) || !nodeIds.has(edge.target))
            throw new Error(`Ребро ${edge.id}: ссылка на несуществующую вершину`);
        if (!Number.isFinite(Number(edge.weight)))
            throw new Error(`Ребро ${edge.id}: некорректный вес`);
    });
    return true;
}
""")

    # ===================================================================
    # 7. РАСПЕЧАТКИ ВХОДНОГО И ВЫХОДНОГО ФАЙЛОВ
    # ===================================================================
    add_heading(doc, "7. Распечатки содержимого входного и выходного файлов программы", 1)
    add_paragraph(
        doc,
        "В качестве примера используется неориентированный граф из четырёх вершин (0, 1, "
        "2, 3). Кратчайший путь ищется от вершины 0 до вершины 3. Входной и выходной "
        "файлы соответствуют одному и тому же запуску программы."
    )
    add_heading(doc, "7.1. Входной файл графа (graph.json)", 2)
    add_code_block(doc, """
{
  "directed": false,
  "nodes": [
    { "id": "n0", "label": "0", "x": 100, "y": 200 },
    { "id": "n1", "label": "1", "x": 300, "y": 120 },
    { "id": "n2", "label": "2", "x": 300, "y": 280 },
    { "id": "n3", "label": "3", "x": 500, "y": 200 }
  ],
  "edges": [
    { "id": "e0", "source": "n0", "target": "n1", "weight": 4 },
    { "id": "e1", "source": "n0", "target": "n2", "weight": 2 },
    { "id": "e2", "source": "n2", "target": "n1", "weight": 1 },
    { "id": "e3", "source": "n1", "target": "n3", "weight": 3 }
  ]
}
""")
    add_heading(doc, "7.2. Выходной файл результата (dijkstra_results.json)", 2)
    add_paragraph(
        doc,
        "Кратчайший путь от вершины 0 до вершины 3: 0 → 2 → 1 → 3 со стоимостью 6 "
        "(2 + 1 + 3). Соответствующий выходной файл:"
    )
    add_code_block(doc, """
{
  "path": ["0", "2", "1", "3"],
  "distance": 6,
  "executionTime": 0.118,
  "operationCount": 24,
  "distances": {
    "0": 0,
    "1": 3,
    "2": 2,
    "3": 6
  },
  "predecessors": {
    "1": "2",
    "2": "0",
    "3": "1"
  }
}
""")

    # ===================================================================
    # 8. СПИСОК ИСТОЧНИКОВ
    # ===================================================================
    add_heading(doc, "8. Список использованных источников", 1)
    add_numbered(doc, [
        "Кормен Т. Х., Лейзерсон Ч. И., Ривест Р. Л., Штайн К. Алгоритмы: построение и "
        "анализ. — 2-е изд. — М.: Вильямс, 2013. — 1296 с.",
        "Ахо А. В., Хопкрофт Д., Ульман Дж. Д. Структуры данных и алгоритмы. — "
        "М.: Вильямс, 2003. — 384 с.",
        "Кнут Д. Э. Искусство программирования. Том 1. Основные алгоритмы. — 3-е изд. — "
        "М.: Вильямс, 2000. — 720 с.",
        "Dijkstra E. W. A note on two problems in connexion with graphs. "
        "Numerische Mathematik, 1959, vol. 1, pp. 269–271.",
        "Cytoscape.js Documentation. — URL: https://js.cytoscape.org/",
        "Chart.js Documentation. — URL: https://www.chartjs.org/docs/latest/",
        "jsPDF Documentation. — URL: https://github.com/parallax/jsPDF",
        "MDN Web Docs: JavaScript, File API, LocalStorage. — "
        "URL: https://developer.mozilla.org/",
    ])

    doc.save(DOCX_PATH)


# ---------------------------------------------------------------------------
# Markdown-версия (для удобного редактирования)
# ---------------------------------------------------------------------------
def build_markdown():
    text = f"""# Пояснительная записка к курсовой работе

Тема: **{COURSEWORK_TITLE}**

Проект: **{PROJECT_TITLE}**

## Содержание

1. Титульный лист.
2. Введение.
3. Краткое описание используемых методов и алгоритмов.
4. Особенности программной реализации алгоритмов.
5. Сравнительный анализ используемых алгоритмов и методов.
6. Листинги программ с обязательными исчерпывающими комментариями.
7. Распечатки содержимого входного и выходного файлов программы.
8. Список использованных литературных и информационных источников.

## 2. Введение

В курсовой работе рассматривается задача поиска кратчайшего пути во взвешенном графе
при неотрицательных весах рёбер. Известные методы её решения: поиск в ширину, алгоритмы
Дейкстры, Беллмана-Форда, Флойда-Уоршелла и A*. По заданию основным выбран алгоритм
Дейкстры. Цель работы — самостоятельная реализация алгоритма и его сравнительный анализ.
Алгоритм и очередь с приоритетами реализованы вручную; библиотеки Cytoscape.js, Chart.js
и jsPDF используются только для визуализации, графиков и отчёта.

## 3. Краткое описание используемых методов и алгоритмов

Граф задаётся множеством вершин V и множеством рёбер E (ориентированные и
неориентированные графы). Основной алгоритм — Дейкстры с очередью приоритетов на основе
бинарной кучи. Шаги:

1. Расстояние до стартовой вершины — 0, до остальных — бесконечность.
2. Стартовая вершина помещается в очередь с приоритетом 0.
3. Извлекается вершина с минимальным расстоянием.
4. Выполняется релаксация соседних рёбер.
5. При улучшении пути обновляются расстояние и предшественник.
6. Путь восстанавливается по массиву предшественников.

Сложность: O((V + E) log V) по времени и O(V + E) по памяти.

## 4. Особенности программной реализации

Основные модули: `graph.js` (структура графа и валидация), `dijkstra.js` (алгоритм и
очередь приоритетов), `visualizer.js` (пошаговая визуализация), `ui.js` (контроллер
интерфейса), `storage.js` (файловый ввод-вывод, PNG, PDF), `tests/unit-tests.js` (тесты).

Основные типы данных: GraphManager, DijkstraAlgorithm, PriorityQueue, а также
`Map` (расстояния и предшественники) и `Set` (посещённые вершины).

Ввод-вывод организован через файлы: вход — JSON-граф (`nodes`, `edges`); выход — JSON/TXT
результата, PNG графа и PDF-отчёт. Перед загрузкой граф проходит строгую проверку
(`validateGraphData`). Разработан набор автоматических тестов (`npm test`).

## 5. Сравнительный анализ

| Метод | Ограничения | Время | Память |
|-------|-------------|-------|--------|
| BFS | невзвешенные графы | O(V + E) | O(V) |
| Дейкстра | без отрицательных весов | O((V + E) log V) | O(V + E) |
| Беллман-Форд | допускает отрицательные веса | O(V·E) | O(V) |
| Флойд-Уоршелл | все пары вершин | O(V³) | O(V²) |
| A* | требует эвристику | зависит от эвристики | O(V) |

Для задачи выбран алгоритм Дейкстры: веса неотрицательны, нужен путь от одной стартовой
вершины, сложность O((V + E) log V) оптимальна для наглядной демонстрации.

## 6. Листинги

Полные исходные файлы — в папке `js`. В DOCX-версии приведены фрагменты очереди с
приоритетами (`PriorityQueue`), основного цикла Дейкстры (`run`) и проверки графа
(`validateGraphData`) с комментариями.

## 7. Распечатки входного и выходного файлов

Пример: неориентированный граф из 4 вершин, путь от 0 до 3.

Входной файл — JSON-граф с массивами `nodes` и `edges`. Выходной файл — JSON-результат
с полями `path` (0 → 2 → 1 → 3), `distance` (6), `distances`, `predecessors`,
`executionTime`, `operationCount`.

## 8. Список использованных источников

1. Кормен Т. Х. и др. Алгоритмы: построение и анализ. — М.: Вильямс, 2013.
2. Ахо А. В., Хопкрофт Д., Ульман Дж. Д. Структуры данных и алгоритмы. — М.: Вильямс, 2003.
3. Кнут Д. Э. Искусство программирования. Том 1. — М.: Вильямс, 2000.
4. Dijkstra E. W. A note on two problems in connexion with graphs. Numerische Mathematik, 1959.
5. Cytoscape.js Documentation: https://js.cytoscape.org/
6. Chart.js Documentation: https://www.chartjs.org/docs/latest/
7. jsPDF Documentation: https://github.com/parallax/jsPDF
8. MDN Web Docs: https://developer.mozilla.org/
"""
    MD_PATH.write_text(text, encoding="utf-8")


def build_readme():
    text = """# Пояснительная записка

В этой папке лежит пояснительная записка к курсовой работе по проекту
«Визуализатор алгоритма Дейкстры».

Файлы:

- `poyasnitelnaya_zapiska.docx` — готовая версия для Word (оформлена по требованиям).
- `poyasnitelnaya_zapiska.md` — текстовая версия, которую удобно редактировать.
- `build_coursework_note.py` — генератор DOCX и Markdown.
- `backup/` — резервные копии предыдущих версий.

Оформление DOCX соответствует требованиям курсовой работы:

- шрифт Times New Roman 14, межстрочный интервал 1,5, основной текст по ширине;
- поля: верхнее 1,5 см, нижнее 1,5 см, левое 2,5 см, правое 1,0 см;
- нумерация листов в правом верхнем углу;
- заголовки — прописными буквами, жирным, по центру;
- таблицы пронумерованы и подписаны, в тексте на них есть ссылки.

В титульном листе оставлены поля-заглушки, которые нужно заменить перед сдачей:

- название образовательной организации;
- факультет / кафедра;
- ФИО студента;
- номер группы;
- ФИО руководителя;
- город.
"""
    README_PATH.write_text(text, encoding="utf-8")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_markdown()
    build_readme()
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {MD_PATH}")
    print(f"Created: {README_PATH}")


if __name__ == "__main__":
    main()
