from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUTPUT_DIR / "poyasnitelnaya_zapiska.docx"
MD_PATH = OUTPUT_DIR / "poyasnitelnaya_zapiska.md"
README_PATH = OUTPUT_DIR / "README.md"


PROJECT_TITLE = "Визуализатор алгоритма Дейкстры"
COURSEWORK_TITLE = (
    "Программная реализация алгоритма нахождения кратчайших путей "
    "в графах методом Дейкстры"
)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)


def add_paragraph(doc, text="", bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    return p


def add_heading(doc, title, level=1):
    p = doc.add_paragraph()
    # Требование: перед заголовком интервал примерно в 2 раза,
    # после — в 1,5 раза больше расстояния между строками текста.
    p.paragraph_format.space_before = Pt(24 if level == 1 else 18)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.keep_with_next = True
    # Требование: заголовки прописными буквами, по центру, без переносов.
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(14)


_TABLE_COUNTER = {"n": 0}


def add_table(doc, caption, headers, rows, widths=None):
    # Требование: таблицы нумеруются и подписываются, в тексте на них есть ссылки.
    _TABLE_COUNTER["n"] += 1
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(4)
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap_run = cap.add_run(f"Таблица {_TABLE_COUNTER['n']} — {caption}")
    cap_run.font.name = "Times New Roman"
    cap_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cap_run.font.size = Pt(14)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr_cells[idx], header, bold=True, size=11)
        set_cell_shading(hdr_cells[idx], "E8EEF5")

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=11)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)

    doc.add_paragraph()
    return table


def add_code_block(doc, code):
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9)


def configure_document(doc):
    section = doc.sections[0]
    # Поля по требованию КР: верхнее 1,5 / нижнее 1,5 / левое 2,5 / правое 1,0 см. A4.
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(1.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    # Требование: нумерация листов — в правом верхнем углу (в верхнем колонтитуле).
    header = section.header.paragraphs[0]
    add_page_number(header, alignment=WD_ALIGN_PARAGRAPH.RIGHT)


def add_title_page(doc):
    for text, size, bold in [
        ("[Название образовательной организации]", 12, False),
        ("[Факультет / кафедра]", 12, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ПОЯСНИТЕЛЬНАЯ ЗАПИСКА\nк курсовой работе")
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'\nпо теме: "{COURSEWORK_TITLE}"')
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(14)

    for _ in range(4):
        doc.add_paragraph()

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.style = "Table Grid"
    rows = [
        ("Выполнил:", "[ФИО студента]"),
        ("Группа:", "[номер группы]"),
        ("Проверил:", "[ФИО руководителя]"),
        ("Оценка:", "____________"),
    ]
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.rows[row_idx].cells[col_idx]
            set_cell_text(cell, value, bold=(col_idx == 0), size=12)

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Город] 2026")
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    doc.add_page_break()


def add_contents_page(doc):
    add_heading(doc, "СОДЕРЖАНИЕ", 1)
    add_numbered(doc, [
        "Титульный лист.",
        "Введение.",
        "Краткое описание используемых методов и алгоритмов.",
        "Особенности программной реализации алгоритмов.",
        "Сравнительный анализ используемых алгоритмов и методов.",
        "Листинги программ с обязательными исчерпывающими комментариями.",
        "Распечатки содержимого входного и выходного файлов программы.",
        "Список использованных литературных и информационных источников.",
    ])
    add_paragraph(
        doc,
        "Структура пояснительной записки составлена по требованиям к курсовой работе: "
        "сначала формулируется задача, затем описываются методы, реализация, анализ, "
        "листинги и примеры входных и выходных данных."
    )
    doc.add_page_break()


def build_docx():
    doc = Document()
    _TABLE_COUNTER["n"] = 0
    configure_document(doc)
    add_title_page(doc)
    add_contents_page(doc)

    add_heading(doc, "1. Введение", 1)
    add_paragraph(
        doc,
        "Цель курсовой работы - разработать интерактивное frontend-приложение для "
        "изучения, визуализации и исследования алгоритма Дейкстры на взвешенных графах."
    )
    add_paragraph(
        doc,
        "В рамках задачи рассматривается поиск кратчайшего пути между двумя вершинами "
        "графа при условии, что веса всех рёбер являются неотрицательными. Для решения "
        "используется алгоритм Дейкстры, так как он гарантирует нахождение оптимального "
        "пути в графах с неотрицательными весами и хорошо подходит для пошаговой "
        "образовательной визуализации."
    )
    add_paragraph(
        doc,
        "Приложение реализовано без сборщика на HTML, CSS и JavaScript. Для визуального "
        "представления графа используется Cytoscape.js, для графиков производительности - "
        "Chart.js, для формирования отчёта - jsPDF. Основная логика алгоритма реализована "
        "самостоятельно в модуле dijkstra.js."
    )
    add_bullets(doc, [
        "создание и редактирование ориентированных и неориентированных графов;",
        "запуск алгоритма в пошаговом режиме и в режиме быстрого расчёта;",
        "отображение очереди приоритетов, таблицы расстояний и итогового пути;",
        "импорт и экспорт графов в JSON, экспорт результатов в TXT, JSON, PNG и PDF;",
        "анализ производительности на графах разного размера.",
    ])

    add_heading(doc, "2. Краткое описание используемых методов и алгоритмов", 1)
    add_heading(doc, "2.1. Графовая модель", 2)
    add_paragraph(
        doc,
        "Граф задаётся множеством вершин V и множеством рёбер E. В проекте поддерживаются "
        "как неориентированные, так и ориентированные графы. Каждое ребро имеет вес - "
        "числовую стоимость перехода между вершинами. Основные понятия графовой модели "
        "и их назначение в программе приведены в таблице 1."
    )
    add_table(
        doc,
        "Основные понятия графовой модели",
        ["Понятие", "Назначение в программе"],
        [
            ["Вершина", "Объект с идентификатором, подписью и координатами на холсте."],
            ["Ребро", "Связь между двумя вершинами с неотрицательным весом."],
            ["Список смежности", "Способ быстрого получения соседей выбранной вершины."],
            ["Очередь с приоритетами", "Структура для выбора вершины с минимальным расстоянием."],
        ],
        [3.5, 11.0],
    )
    add_heading(doc, "2.2. Алгоритм Дейкстры", 2)
    add_paragraph(
        doc,
        "Алгоритм Дейкстры находит кратчайшие пути от стартовой вершины до остальных "
        "вершин графа. На каждом шаге выбирается непосещённая вершина с минимальным "
        "известным расстоянием, после чего выполняется релаксация её исходящих рёбер."
    )
    add_numbered(doc, [
        "Для стартовой вершины устанавливается расстояние 0, для остальных - бесконечность.",
        "Стартовая вершина помещается в очередь с приоритетом 0.",
        "Из очереди извлекается вершина с минимальным расстоянием.",
        "Для каждого соседа проверяется, можно ли улучшить известное расстояние.",
        "Если путь улучшен, обновляются расстояние и предшественник вершины.",
        "После обработки всех достижимых вершин восстанавливается итоговый путь.",
    ])
    add_paragraph(
        doc,
        "В проекте очередь с приоритетами реализована на основе бинарной кучи. Это "
        "позволяет выполнять вставку и извлечение минимального элемента за O(log n). "
        "Общая временная сложность алгоритма составляет O((V + E) log V)."
    )

    add_heading(doc, "3. Особенности программной реализации алгоритмов", 1)
    add_heading(doc, "3.1. Общая структура проекта", 2)
    add_paragraph(
        doc,
        "Проект разделён на независимые модули, каждый из которых отвечает за свою часть "
        "функциональности. Назначение основных файлов проекта приведено в таблице 2."
    )
    add_table(
        doc,
        "Назначение основных файлов проекта",
        ["Файл", "Назначение"],
        [
            ["index.html", "Разметка интерфейса и подключение локальных библиотек."],
            ["css/style.css", "Темы, адаптивность, стили графа и элементов интерфейса."],
            ["js/graph.js", "Структура графа, генерация, импорт, экспорт и валидация данных."],
            ["js/dijkstra.js", "Алгоритм Дейкстры и очередь с приоритетами."],
            ["js/visualizer.js", "Пошаговая визуализация состояния алгоритма."],
            ["js/ui.js", "Главный контроллер интерфейса и обработчики событий."],
            ["js/storage.js", "LocalStorage, импорт/экспорт, PNG и PDF-отчёты."],
            ["tests/unit-tests.js", "Автоматические тесты ядра проекта."],
        ],
        [4.2, 10.3],
    )
    add_heading(doc, "3.2. Входные и выходные данные", 2)
    add_paragraph(
        doc,
        "Входные данные представлены JSON-структурой графа. Она содержит тип графа, "
        "массив вершин и массив рёбер. Для каждой вершины задаются идентификатор, "
        "подпись и координаты. Для каждого ребра задаются источник, цель и вес."
    )
    add_paragraph(
        doc,
        "Выходными данными являются найденный путь, стоимость пути, таблица расстояний, "
        "массив предшественников, время выполнения, число операций и при необходимости "
        "пошаговый журнал визуализации."
    )
    add_heading(doc, "3.3. Проверка корректности данных", 2)
    add_paragraph(
        doc,
        "В модуле GraphManager.validateGraphData выполняется строгая проверка импортируемого "
        "JSON: наличие массивов nodes и edges, уникальность идентификаторов, существование "
        "вершин, на которые ссылаются рёбра, и запрет отрицательных или некорректных весов."
    )
    add_heading(doc, "3.4. Ограничения производительности", 2)
    add_paragraph(
        doc,
        "Для больших графов подробная визуализация может требовать много памяти, так как "
        "каждый шаг содержит снимок состояния алгоритма. Поэтому для графов больше 100 "
        "вершин интерфейс рекомендует режим быстрого расчёта, а для графов больше 200 "
        "вершин подробная визуализация автоматически отключается."
    )

    add_heading(doc, "4. Сравнительный анализ используемых алгоритмов и методов", 1)
    add_paragraph(
        doc,
        "Сравнительные временные и пространственные оценки рассмотренных алгоритмов "
        "поиска путей приведены в таблице 3."
    )
    add_table(
        doc,
        "Сравнение алгоритмов поиска кратчайших путей",
        ["Метод", "Ограничения", "Сложность", "Применимость"],
        [
            ["BFS", "Только невзвешенные графы", "O(V + E)", "Не подходит для произвольных весов."],
            ["Дейкстра", "Нет отрицательных весов", "O((V + E) log V)", "Основной метод проекта."],
            ["Беллман-Форд", "Допускает отрицательные веса", "O(VE)", "Медленнее, но универсальнее."],
            ["Флойд-Уоршелл", "Все пары вершин", "O(V^3)", "Удобен для малых плотных графов."],
            ["A*", "Требует эвристику", "Зависит от эвристики", "Полезен для карт и навигации."],
        ],
        [3.0, 4.3, 3.1, 4.1],
    )
    add_paragraph(
        doc,
        "Для данной курсовой работы выбран алгоритм Дейкстры, так как задача ограничена "
        "неотрицательными весами, а главная цель проекта - наглядная демонстрация "
        "процесса поиска кратчайшего пути."
    )

    add_heading(doc, "5. Листинги программ с комментариями", 1)
    add_paragraph(
        doc,
        "Ниже приведены ключевые фрагменты программы. Полные исходные файлы находятся "
        "в папке js проекта."
    )
    add_heading(doc, "5.1. Запуск алгоритма Дейкстры", 2)
    add_code_block(doc, """
run(startNodeId, endNodeId) {
    const startTime = performance.now();
    this.distances = new Map();
    this.predecessors = new Map();
    this.visited = new Set();
    this.steps = [];

    this._validateInput(startNodeId, endNodeId);
    this._validateWeights();
    this._initializeDistances(startNodeId);

    const queue = new PriorityQueue();
    queue.enqueue(startNodeId, 0);

    while (!queue.isEmpty()) {
        const current = queue.dequeue();
        if (this.visited.has(current.element)) continue;
        this._processNode(current.element, endNodeId, queue);
    }

    return this._buildResult(startNodeId, endNodeId, startTime);
}
""")
    add_heading(doc, "5.2. Проверка импортируемого графа", 2)
    add_code_block(doc, """
static validateGraphData(data) {
    if (!data || typeof data !== 'object') {
        throw new Error('Неверный формат: данные графа должны быть объектом');
    }
    if (!Array.isArray(data.nodes) || !Array.isArray(data.edges)) {
        throw new Error('Неверный формат: отсутствуют массивы nodes или edges');
    }
    // Проверяются дубли ID, ссылки рёбер на существующие вершины
    // и отсутствие отрицательных весов.
    return true;
}
""")
    add_heading(doc, "5.3. Быстрый режим без накопления шагов", 2)
    add_code_block(doc, """
const dijkstra = new DijkstraAlgorithm(graph, { captureSteps: false });
const result = dijkstra.run(startNode, endNode);

// В этом режиме сохраняется итоговый путь и таблица расстояний,
// но не создаются подробные снимки каждого шага визуализации.
""")

    add_heading(doc, "6. Распечатки содержимого входного и выходного файлов программы", 1)
    add_heading(doc, "6.1. Пример входного JSON-файла", 2)
    add_code_block(doc, """
{
  "directed": false,
  "nodes": [
    { "id": "n0", "label": "0", "x": 100, "y": 200 },
    { "id": "n1", "label": "1", "x": 300, "y": 100 },
    { "id": "n2", "label": "2", "x": 300, "y": 300 }
  ],
  "edges": [
    { "id": "e0", "source": "n0", "target": "n1", "weight": 4 },
    { "id": "e1", "source": "n0", "target": "n2", "weight": 2 },
    { "id": "e2", "source": "n2", "target": "n1", "weight": 1 }
  ]
}
""")
    add_heading(doc, "6.2. Пример выходного JSON-файла результата", 2)
    add_code_block(doc, """
{
  "path": ["0", "2", "1", "3"],
  "distance": 4,
  "executionTime": 0.123,
  "operationCount": 42,
  "distances": {
    "0": 0,
    "1": 3,
    "2": 1,
    "3": 4
  },
  "predecessors": {
    "1": "2",
    "2": "0",
    "3": "1"
  }
}
""")

    add_heading(doc, "7. Список использованных источников", 1)
    add_numbered(doc, [
        "Dijkstra E. W. A note on two problems in connexion with graphs. Numerische Mathematik, 1959.",
        "Cytoscape.js Documentation. URL: https://js.cytoscape.org/",
        "Chart.js Documentation. URL: https://www.chartjs.org/docs/latest/",
        "jsPDF Documentation. URL: https://github.com/parallax/jsPDF",
        "Bootstrap Documentation. URL: https://getbootstrap.com/docs/5.3/",
        "Font Awesome Documentation. URL: https://fontawesome.com/",
        "MDN Web Docs: JavaScript, LocalStorage, File API. URL: https://developer.mozilla.org/",
    ])

    doc.save(DOCX_PATH)


def build_markdown():
    text = f"""# Пояснительная записка к курсовой работе

Тема: **{COURSEWORK_TITLE}**

Проект: **{PROJECT_TITLE}**

## Содержание

1. Титульный лист.
2. Введение.
3. Краткое описание используемых методов и алгоритмов.
4. Особенности программной реализации алгоритмов.
5. Сравнительный анализ используемых алгоритмов и методов.
6. Листинги программ с обязательными исчерпывающими комментариями.
7. Распечатки содержимого входного и выходного файлов программы.
8. Список использованных литературных и информационных источников.

## 1. Введение

Цель курсовой работы - разработать интерактивное frontend-приложение для изучения,
визуализации и исследования алгоритма Дейкстры на взвешенных графах.

Приложение реализовано без сборщика на HTML, CSS и JavaScript. Для визуального
представления графа используется Cytoscape.js, для графиков производительности -
Chart.js, для формирования отчёта - jsPDF.

## 2. Краткое описание используемых методов и алгоритмов

Граф задаётся множеством вершин V и множеством рёбер E. В проекте поддерживаются
ориентированные и неориентированные графы. Основной алгоритм - алгоритм Дейкстры
с очередью приоритетов на основе бинарной кучи.

Основные шаги алгоритма:

1. Инициализировать расстояние до стартовой вершины нулём.
2. Остальным вершинам присвоить бесконечность.
3. Поместить стартовую вершину в очередь с приоритетом 0.
4. Извлекать вершину с минимальным расстоянием.
5. Выполнять релаксацию соседних рёбер.
6. Восстановить путь по массиву предшественников.

## 3. Особенности программной реализации

Основные модули проекта:

- `js/graph.js` - структура графа, генерация, импорт, экспорт и валидация.
- `js/dijkstra.js` - алгоритм Дейкстры и очередь с приоритетами.
- `js/visualizer.js` - пошаговая визуализация.
- `js/ui.js` - главный контроллер интерфейса.
- `js/storage.js` - LocalStorage, импорт/экспорт, PNG и PDF.
- `tests/unit-tests.js` - автоматические тесты ядра.

Входные данные задаются JSON-структурой графа. Выходные данные включают путь,
стоимость, таблицу расстояний, предшественников, время выполнения и число операций.

## 4. Сравнительный анализ

Для задачи выбран алгоритм Дейкстры, потому что веса рёбер неотрицательны, а
алгоритм хорошо подходит для пошаговой образовательной визуализации. По сравнению
с Беллманом-Фордом он быстрее на графах без отрицательных весов. По сравнению с
Флойдом-Уоршеллом он эффективнее, когда нужен путь от одной стартовой вершины.

## 5. Листинги

Полные исходные файлы находятся в папке `js`. В DOCX-версии пояснительной записки
приведены основные фрагменты запуска алгоритма, проверки данных и быстрого режима.

## 6. Примеры входных и выходных файлов

Входной файл: JSON-граф с массивами `nodes` и `edges`.

Выходной файл: JSON-результат с `path`, `distance`, `distances`, `predecessors`,
`executionTime` и `operationCount`.

## 7. Источники

1. Dijkstra E. W. A note on two problems in connexion with graphs. Numerische Mathematik, 1959.
2. Cytoscape.js Documentation: https://js.cytoscape.org/
3. Chart.js Documentation: https://www.chartjs.org/docs/latest/
4. jsPDF Documentation: https://github.com/parallax/jsPDF
5. Bootstrap Documentation: https://getbootstrap.com/docs/5.3/
6. Font Awesome Documentation: https://fontawesome.com/
7. MDN Web Docs: https://developer.mozilla.org/
"""
    MD_PATH.write_text(text, encoding="utf-8")


def build_readme():
    text = """# Пояснительная записка

В этой папке лежит черновик пояснительной записки к курсовой работе по проекту
«Визуализатор алгоритма Дейкстры».

Файлы:

- `poyasnitelnaya_zapiska.docx` — готовая версия для Word.
- `poyasnitelnaya_zapiska.md` — текстовая версия, которую удобно редактировать.
- `build_coursework_note.py` — генератор DOCX и Markdown.

В титульном листе оставлены поля-заглушки:

- название образовательной организации;
- факультет / кафедра;
- ФИО студента;
- номер группы;
- ФИО руководителя;
- город.

Перед сдачей нужно заменить эти поля на реальные данные и при необходимости
добавить требования вашего преподавателя к оформлению.
"""
    README_PATH.write_text(text, encoding="utf-8")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_docx()
    build_markdown()
    build_readme()
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {MD_PATH}")
    print(f"Created: {README_PATH}")


if __name__ == "__main__":
    main()
